from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from pathlib import Path
from sys import argv
from abc import ABC, abstractmethod

iso885915_utf_map = {
    ord(u"\u2018"):  ord(u"'"),     # LEFT SINGLE QUOTATION MARK
    ord(u"\u2019"):  ord(u"'"),     # RIGHT SINGLE QUOTATION MARK
    ord(u"\u201c"):  ord(u'"'),     # LEFT DOUBLE QUOTATION MARK
    ord(u"\u201d"):  ord(u'"'),     # RIGHT DOUBLE QUOTATION MARK
    ord(u"\u2026"):  "...",         # HORIZONTAL ELLIPSIS
    ord(u"\u2013"):  ord(u"-"),     # EN DASH
    ord(u"\xe7"):    ord(u"c")      # LATIN SMALL LETTER C WITH CEDILLA
}

class ParagraphProcessor:
    def __init__(self):
        self.buffer = ""
        self.italic_active = False
        self.bold_active = False

    def append(self, text: str):
        self.buffer = self.buffer + text

    def set_italic(self, italic: bool):
        if italic != self.italic_active:
            self.append("*")
        
        self.italic_active = italic

    def set_bold(self, bold: bool):
        if bold != self.bold_active:
            self.append("**")

        self.bold_active = bold

    def append_run(self, run: Run):
        if is_pagebreak(run):
            self.append("\n\\pagebreakNum\n")
            return

        self.set_italic(bool(run.italic))
        self.set_bold(bool(run.bold))
        self.append(run.text.translate(iso885915_utf_map))

    def finalize(self) -> str:
        self.set_bold(False)
        self.set_italic(False)
        return self.buffer

class StyleGroupProcessor(ABC):
    def __init__(self):
        self.initialized = False
        self.lines = []

    @staticmethod
    @abstractmethod
    def accepts_paragraph(paragraph: Paragraph) -> bool:
        return False

    def post_process(self, paragraph_info: Paragraph, paragraph_line: str) -> str:
        return paragraph_line

    def get_paragraph_separator(self):
        return "\n\n"

    def on_finalize(self):
        pass

    def append_line(self, text: str):
        self.lines.append(text)

    def append_paragraph(self, paragraph: Paragraph):
        assert self.accepts_paragraph(paragraph)

        paragraph_processor = ParagraphProcessor()
        for run in paragraph.runs:
            paragraph_processor.append_run(run)

        paragraph_line = paragraph_processor.finalize()
        paragraph_line = self.post_process(paragraph, paragraph_line)
        self.append_line(paragraph_line + self.get_paragraph_separator())

    def finalize(self, output: list):
        self.on_finalize()
        for line in self.lines:
            output.append(line)

class HeaderGroupProcessor(StyleGroupProcessor):
    @staticmethod
    def accepts_paragraph(paragraph: Paragraph) -> bool:
        return paragraph.style.name.lower().startswith("heading")

    def post_process(self, paragraph_info: Paragraph, paragraph_line: str) -> str:
        header_level = int(paragraph_info.style.name.split(" ")[1])
        header_prefix = "#" * header_level + " "
        return header_prefix + paragraph_line

class DescriptiveGroupProcessor(StyleGroupProcessor):
    descriptive_header_added = False
    
    @staticmethod
    def accepts_paragraph(paragraph: Paragraph) -> bool:
        return paragraph.style.name.lower().startswith("descriptive")

    def post_process(self, paragraph_info: Paragraph, paragraph_line: str) -> str:
        if not self.descriptive_header_added:
            self.add_descriptive_header()

        return paragraph_line

    def on_finalize(self):
        if self.descriptive_header_added:
            self.close_descriptive_header()

    def add_descriptive_header(self):
        self.append_line('\n\n<div class="descriptive">\n\n')
        self.descriptive_header_added = True

    def close_descriptive_header(self):
        self.append_line("\n\n</div>\n\n")

class GreenBoxGroupProcessor(StyleGroupProcessor):
    HEADER_PREFIX = "#" * 5 + " "
    first_paragraph = True

    @staticmethod
    def accepts_paragraph(paragraph: Paragraph) -> bool:
        return paragraph.style.name.lower() == "green box"

    def post_process(self, paragraph_info: Paragraph, paragraph_line: str) -> str:
        prefix = "> "
        if self.first_paragraph:
            prefix = prefix + self.HEADER_PREFIX
            self.first_paragraph = False

        return prefix + paragraph_line

    def get_paragraph_separator(self):
        return "\n>\n"

class ListGroupProcessor(StyleGroupProcessor):
    @staticmethod
    def accepts_paragraph(paragraph: Paragraph) -> bool:
        return paragraph.style.name.lower().startswith("list")

    def post_process(self, paragraph_info: Paragraph, paragraph_line: str) -> str:
        return "* " + paragraph_line

    def get_paragraph_separator(self):
        return "\n"

class NormalGroupProcessor(StyleGroupProcessor):
    @staticmethod
    def accepts_paragraph(paragraph: Paragraph) -> bool:
        return paragraph.style.name.lower() == "normal"

def is_pagebreak(run: Run) -> bool:
    return 'w:br' in run._element.xml and 'type="page"' in run._element.xml

def make_group_processor_for_paragraph(paragraph: Paragraph) -> StyleGroupProcessor:
    if HeaderGroupProcessor.accepts_paragraph(paragraph):
        return HeaderGroupProcessor()

    elif DescriptiveGroupProcessor.accepts_paragraph(paragraph):
        return DescriptiveGroupProcessor()

    elif GreenBoxGroupProcessor.accepts_paragraph(paragraph):
        return GreenBoxGroupProcessor()

    elif ListGroupProcessor.accepts_paragraph(paragraph):
        return ListGroupProcessor()

    elif NormalGroupProcessor.accepts_paragraph(paragraph):
        return NormalGroupProcessor()

    else:
        raise Exception("Can't find processor for style " + paragraph.style.name)

def convert_document(doc: Document, out):
    if len(doc.paragraphs) == 0:
        print("Document has not paragraphs.")
        return

    output_lines = []
    style_group_processor = make_group_processor_for_paragraph(doc.paragraphs[0])
    assert style_group_processor is not None

    for paragraph in doc.paragraphs:
        if not style_group_processor.accepts_paragraph(paragraph):
            style_group_processor.finalize(output_lines)            
            style_group_processor = make_group_processor_for_paragraph(paragraph)
            assert style_group_processor is not None

        style_group_processor.append_paragraph(paragraph)
        
    style_group_processor.finalize(output_lines)
    out.writelines(output_lines)

# ============================
# MAIN
# ============================

if len(argv) <= 2:
    print("Usage: python docx-to-md.py <source-file> <out-file>")
    exit(1)

source_path = Path(argv[1])
if not source_path.is_file():
    print(source_path.name)
    print("Source file is not a file.")

source_document = Document(source_path.resolve())
out_path = argv[2]
with open(out_path, 'w') as out_file:
    convert_document(source_document, out_file)